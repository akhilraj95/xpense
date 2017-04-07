import secret
from models import User,Currency,Expense,Bill
import urllib2,json
import urllib2
import urllib
from docx import Document
import time
from docx.shared import Inches
import os
from datetime import timedelta

from django.core.files import File

#FlockOS
from pyflock import FlockClient, verify_event_token
from pyflock import Message, SendAs, Attachment, Views, WidgetView, HtmlView, ImageView, Image, Download, Button, OpenWidgetAction, OpenBrowserAction, SendToAppAction


def appInstall(pjson):
    try:
        userId = pjson['userId']
        token = pjson['token']
        u = User(userId = userId, token=token)
        u.save()
    except:
        raise

def appUninstall(pjson):
    try:
        userId = pjson['userId']
        User.objects.get(userId=userId).delete()
    except:
        raise

def sendMessage(chat_id,userId,message):
    try:
        user = User.objects.get(userId=userId)
        flock_client = FlockClient(token=user.token, app_id=secret.getAppID)
        send_as_xpense = SendAs(name='Xpense', profile_image='https://pbs.twimg.com/profile_images/1788506913/HAL-MC2_400x400.png')
        send_as_message = Message(to=chat_id,text=message,send_as=send_as_xpense)
        flock_client.send_chat(send_as_message)
    except:
        raise

def total(expense_list,trackObj):
    try:
        currency_list = Currency.objects.all()
        total = {}
        for curr in currency_list:
            total[curr.abbr] = 0
        for expense in expense_list:
            total[expense.currency.abbr]+=expense.amount
        total = {key: value for key, value in total.items()
             if value is not 0}
        target_curr = str(trackObj.budget_currency.abbr)
        target_value = 0
        perc_spent = 0
        if target_curr in total.keys():
            target_value = total[target_curr]

        #check if there is a budget
        if trackObj.budget != 0:
            #converting all currencies to budget currency
            target_value = 0
            symbols = total.keys()
            rates = getconversionrates(target_curr,symbols)
            target_value = 0
            for key, value in total.items():
                if key==target_curr:
                    target_value+=value
                else:
                    target_value+= round(value/rates[key],2)

            perc_spent = round((target_value/trackObj.budget)*100,2)
        return total,target_value,perc_spent
    except:
        raise

def getconversionrates(to_curr,from_curr_list):
    if to_curr in from_curr_list:
        from_curr_list.remove(to_curr)
    from_curr_list_str = ','.join(from_curr_list)
    API_string = 'http://api.fixer.io/latest?base='+to_curr+'&symbols='+from_curr_list_str
    response = urllib2.urlopen(API_string).read()
    pjson = json.loads(response)
    return pjson['rates']

def fetchMessagePictures(group_id,token,uids):
    data = [('chat',str(group_id)),('token',str(token)),('uids',uids)]
    url = 'https://api.flock.co/v1/chat.fetchMessages'
    req = urllib2.Request(url, headers={'Content-Type' : 'application/x-www-form-urlencoded'})
    result = urllib2.urlopen(req, urllib.urlencode(data))
    content = result.read()
    content = json.loads(content)
    src_list = []
    for data in content:
        for attachment in data['attachments']:
            for fil in attachment['downloads']:
                if str(fil['mime']) in ['image/jpeg','image/png']:
                    src_list.append(fil['src'])
    return src_list


def report(track,userId):
    document = Document()
    document.add_heading('Expense Report - '+ str(track.name))

    status = '\nPurpose : '+str(track.purpose)+'\n'

    now = time.strftime("%c")

    status = status + 'Report date: '+str(now)+'\n'
    expense_list = Expense.objects.filter(track = track)

    user = User.objects.get(userId=userId)
    flock_client = FlockClient(token=user.token, app_id=secret.getAppID)
    pjson = flock_client.get_user_info()
    utc = str(pjson['timezone'])
    hours = int(utc[1:3])
    minutes = int(utc[4:6])
    if(utc[0]=='+'):
        for expense in expense_list:
            expense.timestamp += timedelta(hours=hours,minutes=minutes)
    else:
        for expense in expense_list:
            expense.timestamp -= timedelta(hours=hours,minutes=minutes)

    total_expense_by_curr,converted_total,perc_spent = total(expense_list,track)

    if track.budget != 0:
        status = status + 'Budget: '+str(track.budget_currency.abbr)+' '+str(track.budget)+'\n'
        status = status + 'Total spent: '+str(track.budget_currency.abbr)+' '+ str(converted_total)+'\n'
        status = status + 'Spending: '+str(perc_spent)+'%'+'\n'

    status = status + 'Spending per currency:\n'
    for key,value in total_expense_by_curr.items():
        status = status+' - '+str(key)+' '+str(value)+'\n'

    paragraph = document.add_paragraph(status)

    #table
    table = document.add_table(rows=1, cols=4)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Purpose'
    heading_cells[1].text = 'Paid By'
    heading_cells[2].text = 'Time'
    heading_cells[3].text = 'Amount'

    for expense in expense_list:
        cells = table.add_row().cells
        cells[0].text = expense.purpose
        cells[1].text = expense.paidby
        cells[2].text = str(expense.timestamp)
        cells[3].text = expense.currency.abbr +' '+str(expense.amount)

    filename = 'media/'+str(track.chat.chat_id[2:])+'_'+str(track.id)+'.docx'
    download_bills(expense_list,document)

    document.save(filename)
    django_file = File(open(filename,'r'))
    print (django_file.name)
    return django_file



def download_bills(expense_list,document):
    url_list = []
    for expense in expense_list:
        ul = Bill.objects.filter(expense = expense)
        for u in ul:
            url_list.append(str(u.url))
    print url_list
    if(len(url_list)):
        document.add_page_break()
        file_list = []
        for url in url_list:
            f = open(url[33:],'wb')
            f.write(urllib.urlopen(url).read())
            f.close()
            file_list.append(url[33:])

        for fil in file_list:
            document.add_picture(fil,width=Inches(6.0))
            os.remove(fil)
